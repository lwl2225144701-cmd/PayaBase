import asyncio
import json
import uuid
import logging
import tempfile
import os
from typing import Optional

from sqlalchemy import text
from sqlalchemy.ext.asyncio import create_async_engine

from core.config import settings
from core.embedding.client import SyncEmbeddingClient
from models.tables import Document, Chunk
from minio import Minio

logger = logging.getLogger(__name__)


def index_document(document_id: str):
    asyncio.run(_index_async(document_id))


async def _index_async(document_id: str):
    engine = create_async_engine(settings.database_url)
    async with engine.begin() as conn:
        doc_result = await conn.execute(
            text("SELECT id, file_path, file_type, title, knowledge_base_id FROM documents WHERE id=:id"),
            {"id": document_id}
        )
        doc_row = doc_result.mappings().first()
        if not doc_row:
            return

        try:
            await conn.execute(
                text("UPDATE documents SET status='indexing' WHERE id=:id"),
                {"id": document_id}
            )

            minio_client = Minio(
                settings.minio_endpoint,
                access_key=settings.minio_access_key,
                secret_key=settings.minio_secret_key,
                secure=False,
            )
            response = minio_client.get_object(settings.minio_bucket, doc_row["file_path"])
            file_data = response.read()
            file_ext = doc_row["file_type"].lower()

            text_content = _extract_text(file_data, file_ext)
            chunks_data = []

            if text_content.strip():
                words = text_content.split()
                chunk_size = 500
                for i in range(0, len(words), chunk_size):
                    chunk_text = " ".join(words[i : i + chunk_size])
                    chunks_data.append(
                        {
                            "content": chunk_text,
                            "meta": {"source": doc_row["title"]},
                            "token_count": len(chunk_text.split()),
                        }
                    )

            # Delete old chunks before reindex
            await conn.execute(
                text("DELETE FROM chunks WHERE document_id=:doc_id"),
                {"doc_id": str(doc_row["id"])}
            )

            # Generate embeddings
            texts = [c["content"] for c in chunks_data]
            embedding_client = SyncEmbeddingClient()
            vectors = embedding_client.embed_batch(texts)

            # Create chunks
            for chunk_data, vector in zip(chunks_data, vectors):
                chunk_id = str(uuid.uuid4())
                await conn.execute(
                    text("INSERT INTO chunks (id, document_id, content, vector, meta, token_count) VALUES (:id, :doc_id, :content, :vector, :meta, :token_count)"),
                    {
                        "id": chunk_id,
                        "doc_id": str(doc_row["id"]),
                        "content": chunk_data["content"],
                        "vector": vector,  # 直接存List[float]
                        "meta": json.dumps(chunk_data["meta"]),
                        "token_count": chunk_data["token_count"],
                    }
                )

            await conn.execute(
                text("UPDATE documents SET status='ready', chunk_count=:count, indexed_at=NOW() WHERE id=:id"),
                {"id": document_id, "count": len(chunks_data)}
            )
            logger.info(f"Indexed {len(chunks_data)} chunks for document {document_id}")

        except Exception as e:
            await conn.execute(
                text("UPDATE documents SET status='error' WHERE id=:id"),
                {"id": document_id}
            )
            logger.error(f"Indexing failed for {document_id}: {e}")
            raise e

    await engine.dispose()


def _extract_text(file_data: bytes, file_ext: str) -> str:
    if file_ext == "pdf":
        import io
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_data))
        parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
        return "\n".join(parts)
    elif file_ext in ("md", "txt"):
        return file_data.decode("utf-8", errors="replace")
    elif file_ext in ("docx",):
        from docx import Document as DocxDocument
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(file_data)
            tmp_path = tmp.name
        try:
            doc = DocxDocument(tmp_path)
            return "\n".join([p.text for p in doc.paragraphs])
        finally:
            os.unlink(tmp_path)
    else:
        return file_data.decode("utf-8", errors="replace")


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i : i + chunk_size])
        if chunk:
            chunks.append(chunk)
    return chunks