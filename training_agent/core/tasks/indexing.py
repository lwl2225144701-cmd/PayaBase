"""Celery Indexing Task.

5阶段索引管道：
- Stage1: 下载MinIO(0-10%)
- Stage2: 解析+清洗(10-30%)
- Stage3: 语义分块(30-80%)
- Stage4: 向量化(30-80%)
- Stage5: 批量入库(80-100%)
"""

import base64
import io
import json
import logging
import uuid
import time
from datetime import datetime
from typing import Optional

from sqlalchemy import create_engine, text
from sqlalchemy.orm import Session
from langchain_core.documents import Document as LCDocument

from core.config import settings
from core.tasks import celery_app
from models.tables import Chunk, Document

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 重试配置
MAX_RETRIES = 3
RETRY_DELAYS = [60, 300, 600]  # 指数退避: 1分钟, 5分钟, 10分钟


def update_document_status(
    document_id: str,
    status: str,
    progress: int = 0,
    chunk_count: int = 0,
    error_message: Optional[str] = None,
):
    """更新文档状态"""
    engine = create_engine(settings.sync_database_url)
    with Session(engine) as db:
        db.execute(
            text("""
                UPDATE documents 
                SET status=:status, progress=:progress, 
                    chunk_count=:chunk_count, error_message=:error,
                    indexed_at=CASE WHEN :status='ready' THEN NOW() ELSE indexed_at END
                WHERE id=:id
            """),
            {
                "id": document_id,
                "status": status,
                "progress": progress,
                "chunk_count": chunk_count,
                "error": error_message,
            }
        )
        db.commit()
    engine.dispose()


def download_from_minio(file_path: str) -> bytes:
    """Stage1: 从MinIO下载文件"""
    from core.infrastructure.minio.client import get_minio_client
    
    minio_client = get_minio_client()
    
    response = minio_client.get_object(settings.minio_bucket, file_path)
    return response.read()


def extract_text(file_data: bytes, file_ext: str) -> str:
    """Stage2: 解析文档为文本"""
    import io
    import tempfile
    import os
    
    if file_ext == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_data))
        parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
        return "\n".join(parts)
    elif file_ext in ("md", "txt"):
        return file_data.decode("utf-8", errors="replace")
    elif file_ext in ("docx",):
        from docx import Document as DocxDocument
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(file_data)
            tmp_path = tmp.name
        try:
            doc = DocxDocument(tmp_path)
            return "\n".join([p.text for p in doc.paragraphs])
        finally:
            os.unlink(tmp_path)
    else:
        return file_data.decode("utf-8", errors="replace")


IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "gif", "webp", "bmp"}
WORD_EXTENSIONS = {"doc", "docx"}


def _base_metadata(
    kb_id: str,
    doc_id: str,
    tenant_id: str,
    file_ext: str,
    parser_type: str,
    title: str = "",
    source_type: str = "local",
    source_url: str | None = None,
) -> dict:
    return {
        "kb_id": kb_id,
        "doc_id": doc_id,
        "tenant_id": tenant_id,
        "file_type": file_ext,
        "parser_type": parser_type,
        "title": title,
        "source_type": source_type or "local",
        "source_url": source_url,
    }


def parse_markdown_document(
    file_data: bytes,
    kb_id: str,
    doc_id: str,
    tenant_id: str,
    file_ext: str,
    title: str = "",
    source_type: str = "local",
    source_url: str | None = None,
) -> list[LCDocument]:
    """Parse Markdown/text while preserving structure for markdown-aware chunking."""
    text_content = file_data.decode("utf-8", errors="replace")
    return [
        LCDocument(
            page_content=text_content,
            metadata=_base_metadata(kb_id, doc_id, tenant_id, file_ext, "markdown", title, source_type, source_url),
        )
    ]


def parse_word_document(
    file_data: bytes,
    kb_id: str,
    doc_id: str,
    tenant_id: str,
    file_ext: str,
    title: str = "",
    source_type: str = "local",
    source_url: str | None = None,
) -> list[LCDocument]:
    """Parse Word documents, including table cells for retrieval."""
    if file_ext == "doc":
        raise ValueError("暂不支持旧版 .doc 二进制格式，请转换为 .docx 后上传")

    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(file_data))
    parts: list[str] = []

    for para in doc.paragraphs:
        text_value = para.text.strip()
        if text_value:
            parts.append(text_value)

    for table_index, table in enumerate(doc.tables, 1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"[表格{table_index}]\n" + "\n".join(rows))

    text_content = "\n\n".join(parts)
    return [
        LCDocument(
            page_content=text_content,
            metadata=_base_metadata(kb_id, doc_id, tenant_id, file_ext, "word", title, source_type, source_url),
        )
    ]


def parse_image_document(
    file_data: bytes,
    kb_id: str,
    doc_id: str,
    tenant_id: str,
    file_ext: str,
    title: str = "",
    source_type: str = "local",
    source_url: str | None = None,
) -> list[LCDocument]:
    """Parse image documents into retrievable text with Vision LLM."""
    if not settings.index_enable_image_vision:
        raise ValueError("图片解析未启用，请设置 INDEX_ENABLE_IMAGE_VISION=true")

    from core.llm.factory import get_llm_client, is_vision_enabled

    if not is_vision_enabled():
        raise ValueError("图片解析需要配置 LLM_VISION_MODEL")
    from core.prompts.vision import VISION_PROMPT

    mime = f"image/{'jpeg' if file_ext == 'jpg' else file_ext}"
    image_b64 = base64.b64encode(file_data).decode("utf-8")
    result = get_llm_client("vision").chat_with_image(image_b64, VISION_PROMPT, mime_type=mime)
    text_content = f"【图片解析结果】\n{result.strip()}"

    return [
        LCDocument(
            page_content=text_content,
            metadata=_base_metadata(kb_id, doc_id, tenant_id, file_ext, "image_vision", title, source_type, source_url),
        )
    ]


def parse_document_to_langchain(
    file_data: bytes,
    file_ext: str,
    kb_id: str,
    doc_id: str,
    tenant_id: str,
    title: str = "",
    source_type: str = "local",
    source_url: str | None = None,
) -> list[LCDocument]:
    """Unified parser adapter for the indexing pipeline."""
    if file_ext == "pdf":
        from core.rag.pdf_parser import PDFParser

        parser = PDFParser()
        documents = parser.extract(
            file_data,
            kb_id=kb_id,
            doc_id=doc_id,
            tenant_id=tenant_id,
        )
        for doc in documents:
            doc.metadata["file_type"] = file_ext
            doc.metadata["parser_type"] = "pdf"
            doc.metadata["title"] = title
            doc.metadata["source_type"] = source_type or "local"
            doc.metadata["source_url"] = source_url
        return documents

    if file_ext in ("md", "txt"):
        return parse_markdown_document(file_data, kb_id, doc_id, tenant_id, file_ext, title, source_type, source_url)

    if file_ext in WORD_EXTENSIONS:
        return parse_word_document(file_data, kb_id, doc_id, tenant_id, file_ext, title, source_type, source_url)

    if file_ext in IMAGE_EXTENSIONS:
        return parse_image_document(file_data, kb_id, doc_id, tenant_id, file_ext, title, source_type, source_url)

    text_content = extract_text(file_data, file_ext)
    return [
        LCDocument(
            page_content=text_content,
            metadata=_base_metadata(kb_id, doc_id, tenant_id, file_ext, "plain_text", title, source_type, source_url),
        )
    ]


def clean_text(text: str) -> str:
    """Stage2: 脏数据清洗"""
    import re
    lines = text.split("\n")
    cleaned = []
    
    for line in lines:
        # 去除空行
        if not line.strip():
            continue
        # 去除纯数字行
        if line.strip().isdigit():
            continue
        # 去除纯符号行
        if all(c in " \t.-_+=*#@$%^&()[]{}|,.<>/?~`" for c in line.strip()):
            continue
        # 去除过短的行
        if len(line.strip()) < 10 and len(line.strip()) > 0:
            continue
        # 去除页码（如 "Page 1 of 10"）
        if re.match(r'^Page\s+\d+\s+of\s+\d+$', line.strip()):
            continue
        # 去除脚注编号（如 [1], [2]）
        if re.match(r'^\[\d+\]$', line.strip()):
            continue
        cleaned.append(line)
    
    return "\n".join(cleaned)


def split_by_headers(text: str) -> list[str]:
    """按一级标题分割 (# 章节标题)"""
    import re
    # 只按一级标题#分割，不按##或###分割
    parts = re.split(r'\n(?=^#{1}\s+)', text, flags=re.MULTILINE)
    return [p.strip() for p in parts if p.strip()]


def split_by_paragraphs(text: str) -> list[str]:
    """按段落分割"""
    paras = text.split('\n\n')
    return [p.strip() for p in paras if p.strip()]


def split_by_sentences(text: str) -> list[str]:
    """按句子分割（。！？）"""
    import re
    # 按句子结束符分割
    sentences = re.split(r'([。！？\.!?]+)', text)
    result = []
    for i in range(0, len(sentences) - 1, 2):
        sent = sentences[i].strip()
        if sent:
            # 合并句子和结束符
            end = sentences[i + 1] if i + 1 < len(sentences) else ""
            result.append(sent + end)
    return [s for s in result if s.strip()]


def recursive_chunk(text: str, chunk_size: int = 500) -> list[dict]:
    """策略1: 递归字符切分 - 使用字符数处理中文"""
    max_chunk_chars = chunk_size * 3
    min_chunk_chars = int(max_chunk_chars * 0.5)
    
    sections = split_by_headers(text)
    if len(sections) > 1:
        logger.info(f"[Recursive] 按标题分割为 {len(sections)} 个章节")
    
    chunks = []
    current_text = ""
    
    def get_char_size(t: str) -> int:
        return len(t)
    
    for section in sections:
        paragraphs = split_by_paragraphs(section)
        
        for para in paragraphs:
            para_size = get_char_size(para)
            current_size = get_char_size(current_text)
            
            if para_size > max_chunk_chars:
                if current_text.strip():
                    chunks.append(current_text.strip())
                    current_text = ""
                sentences = split_by_sentences(para)
                for sent in sentences:
                    sent_size = get_char_size(sent)
                    if sent_size > max_chunk_chars:
                        chunks.append(sent[:max_chunk_chars])
                        continue
                    if get_char_size(current_text) + sent_size > max_chunk_chars:
                        if current_text.strip():
                            chunks.append(current_text.strip())
                        current_text = sent
                    else:
                        current_text += sent
                continue
            
            if current_size >= min_chunk_chars and current_size + para_size > max_chunk_chars:
                if current_text.strip():
                    chunks.append(current_text.strip())
                    overlap_chars = current_text[-int(len(current_text) * 0.15):]
                    current_text = overlap_chars + "\n\n" + para
                else:
                    current_text = para
            else:
                if current_text:
                    current_text += "\n\n" + para
                else:
                    current_text = para
    
    if current_text.strip():
        chunks.append(current_text.strip())
    
    result = []
    for chunk in chunks:
        char_count = get_char_size(chunk)
        result.append({
            "content": chunk[:10000],
            "meta": {"source": "recursive", "size": char_count},
            "token_count": char_count // 3 if char_count > 0 else 0,
        })
    
    logger.info(f"[Recursive] 分块完成: {len(result)} 个chunks")
    return result


def semantic_chunk_by_embedding(text: str, chunk_size: int = 500, similarity_threshold: float = 0.3) -> list[dict]:
    """策略2: 语义分块 - 按Embedding相似度切分
    
    计算相邻段落的embedding相似度，当相似度低于阈值时才切分
    保证每个块在讨论同一个核心概念
    """
    from core.embedding.client import SyncEmbeddingClient
    
    embedding_client = SyncEmbeddingClient()
    max_chars = chunk_size * 3
    
    paragraphs = split_by_paragraphs(text)
    
    if not paragraphs:
        return [{"content": text[:10000], "meta": {"source": "semantic"}, "token_count": len(text) // 3}]
    
    vectors = embedding_client.embed_batch(paragraphs)
    
    chunks = []
    current_chunk = paragraphs[0]
    current_vector = vectors[0]
    
    for i in range(1, len(paragraphs)):
        para = paragraphs[i]
        vec = vectors[i]
        
        sim = cosine_similarity(current_vector, vec)
        current_chars = len(current_chunk)
        new_chars = current_chars + len(para)
        
        if sim < similarity_threshold or new_chars > max_chars * 1.5:
            chunks.append({
                "content": current_chunk[:10000],
                "meta": {"source": "semantic", "size": current_chars, "similarity": round(sim, 3)},
                "token_count": current_chars // 3,
            })
            
            overlap_chars = current_chunk[-int(len(current_chunk) * 0.15):]
            current_chunk = overlap_chars + "\n\n" + para
            current_vector = vec
        else:
            current_chunk += "\n\n" + para
    
    if current_chunk.strip():
        chunks.append({
            "content": current_chunk[:10000],
            "meta": {"source": "semantic", "size": len(current_chunk)},
            "token_count": len(current_chunk) // 3,
        })
    
    logger.info(f"[Semantic] 语义分块完成: {len(chunks)} 个chunks")
    return chunks




def cosine_similarity(a: list[float], b: list[float]) -> float:
    """计算余弦相似度"""
    if not a or not b or len(a) != len(b):
        return 0.0
    
    dot = sum(x * y for x, y in zip(a, b))
    norm_a = sum(x * x for x in a) ** 0.5
    norm_b = sum(x * x for x in b) ** 0.5
    
    if norm_a == 0 or norm_b == 0:
        return 0.0
    
    return dot / (norm_a * norm_b)


def semantic_chunk(text: str, chunk_size: int = 500, overlap: int = 75) -> list[dict]:
    """Stage3: 语义分块 (按段落 + 15%重叠) - 默认使用递归切分"""
    # 优先使用递归字符切分，保留原逻辑作为fallback
    return recursive_chunk(text, chunk_size)


def embed_texts(texts: list[str]) -> list[list[float]]:
    """Stage4: 向量化"""
    from core.embedding.client import SyncEmbeddingClient
    
    embedding_client = SyncEmbeddingClient()
    
    # 批量向量化
    vectors = embedding_client.embed_batch(texts)
    
    return vectors


def get_chunk_policy(file_ext: str) -> tuple[int, int, str]:
    if file_ext == "md":
        return settings.index_md_chunk_size, settings.index_chunk_overlap, "md_structured"
    if file_ext in IMAGE_EXTENSIONS:
        return settings.index_image_chunk_size, min(50, settings.index_chunk_overlap), "image_vision"
    if file_ext in WORD_EXTENSIONS:
        return settings.index_word_chunk_size, settings.index_chunk_overlap, "word_structured"
    return 500, 50, "recursive"


def build_chunks_data(documents: list[LCDocument], file_ext: str) -> list[dict]:
    """Build chunk records with file-type aware chunking metadata."""
    from core.rag.chunker import TextChunker

    chunk_size, chunk_overlap, chunk_type = get_chunk_policy(file_ext)
    chunker = TextChunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    chunk_docs = chunker.chunk_documents(documents, file_type=file_ext)

    chunks_data = []
    for chunk_doc in chunk_docs:
        content = chunk_doc.page_content.strip()
        if not content:
            continue

        metadata = dict(chunk_doc.metadata or {})
        metadata.update(
            {
                "chunk_strategy": chunk_type,
                "chunk_size": chunk_size,
                "chunk_overlap": chunk_overlap,
            }
        )
        chunks_data.append(
            {
                "content": content,
                "chunk_id": str(uuid.uuid4()),
                "token_count": metadata.get("token_count", len(content) // 3),
                "chunk_type": chunk_type,
                "meta": metadata,
            }
        )

    return chunks_data


def build_embedding_text(chunk: dict) -> str:
    """Use the raw chunk content for embedding (索引期不再生成摘要)。

    直接对原文做向量化,索引期零 LLM 调用;HyDE 召回增强改在查询时完成。
    """
    content = (chunk.get("content") or "").strip()
    chunk_type = chunk.get("chunk_type", "")

    if chunk_type == "image_vision":
        return f"图片内容检索文本：{content}"
    if chunk_type == "md_structured":
        return f"Markdown文档片段：{content}"
    if chunk_type == "word_structured":
        return f"Word文档片段：{content}"
    return content


def batch_insert_chunks(
    document_id: str,
    chunks_data: list[dict],
    vectors: list[list[float]],
):
    """Stage5: 批量入库（事务控制：chunks + 状态更新在同一事务）"""
    from datetime import datetime

    if len(chunks_data) != len(vectors):
        raise ValueError(f"向量数量不匹配: chunks={len(chunks_data)}, vectors={len(vectors)}")
    
    engine = create_engine(settings.sync_database_url, pool_pre_ping=True)
    
    try:
        with engine.begin() as conn:
            # 0. 先删除该文档的旧 chunk（重索引时必须清理，否则新旧向量并存导致检索混乱）
            conn.execute(
                text("DELETE FROM chunks WHERE document_id=:doc_id"),
                {"doc_id": document_id},
            )

            # 1. 批量插入 chunks
            for chunk_data, vector in zip(chunks_data, vectors):
                chunk_id = chunk_data.get("chunk_id") or str(uuid.uuid4())
                
                conn.execute(
                    text("""
                        INSERT INTO chunks (id, document_id, content, summary, hypothetical_questions, chunk_type, vector, meta, token_count)
                        VALUES (:id, :doc_id, :content, :summary, :hyde, :chunk_type, :vector, :meta, :token_count)
                    """),
                    {
                        "id": chunk_id,
                        "doc_id": document_id,
                        "content": chunk_data["content"],
                        "summary": chunk_data.get("summary", "")[:500],
                        "hyde": json.dumps(chunk_data.get("hypothetical_questions", [])),
                        "chunk_type": chunk_data.get("chunk_type", "recursive"),
                        "vector": vector,
                        "meta": json.dumps(chunk_data.get("meta", {})),
                        "token_count": chunk_data.get("token_count", 0),
                    }
                )
            
            # 2. 更新文档状态为 ready（在同一事务中）
            conn.execute(
                text("""
                    UPDATE documents 
                    SET status=:status, progress=100, chunk_count=:chunk_count, indexed_at=:indexed_at
                    WHERE id=:id
                """),
                {
                    "id": document_id,
                    "status": "ready",
                    "chunk_count": len(chunks_data),
                    "indexed_at": datetime.utcnow(),
                }
            )
            
            # 事务自动提交，全部成功
    except Exception as e:
        # 事务自动回滚
        raise
    finally:
        engine.dispose()


@celery_app.task(bind=True, max_retries=MAX_RETRIES, time_limit=3600, soft_time_limit=3000)
def index_document_task(self, document_id: str):
    """Celery任务: 文档索引5阶段管道
    
    Args:
        document_id: 文档ID
    
    Returns:
        任务结果
    """
    start_time = time.time()
    logger.info(f"[Celery] 开始索引文档: {document_id}")
    
    # 创建数据库引擎
    engine = create_engine(settings.sync_database_url)
    
    try:
        # ===== Stage1: 获取文档信息 =====
        with Session(engine) as db:
            result = db.execute(
                text("""
                    SELECT d.id, d.file_path, d.file_type, d.title, d.source_type, d.source_url, d.knowledge_base_id, kb.tenant_id 
                    FROM documents d 
                    JOIN knowledge_bases kb ON d.knowledge_base_id = kb.id 
                    WHERE d.id=:id
                """),
                {"id": document_id}
            )
            doc_row = result.mappings().first()
            
            if not doc_row:
                logger.error(f"[Celery] 文档不存在: {document_id}")
                return {"status": "error", "message": "Document not found"}
            
            file_path = doc_row["file_path"]
            file_ext = doc_row["file_type"]
            title = doc_row["title"]
            source_type = doc_row["source_type"] or "local"
            source_url = doc_row["source_url"]
            kb_id = str(doc_row["knowledge_base_id"])
            tenant_id = str(doc_row["tenant_id"])
        
        # 更新状态为indexing
        update_document_status(document_id, "indexing", progress=5)
        
        # ===== Stage2: 下载MinIO =====
        logger.info(f"[Stage1] 开始下载文件: {file_path}")
        try:
            file_data = download_from_minio(file_path)
            logger.info(f"[Stage1] 下载成功, 文件大小: {len(file_data)} bytes")
        except Exception as e:
            logger.error(f"[Stage1] 下载失败: {e}")
            update_document_status(document_id, "failed", error_message=f"下载失败: {str(e)}")
            raise
        
        update_document_status(document_id, "indexing", progress=10)
        
        # ===== Stage3: 解析文档 =====
        logger.info(f"[Stage2] 开始解析文档, 类型: {file_ext}")
        
        from core.rag.cleaner import CleanProcessor

        try:
            documents = parse_document_to_langchain(
                file_data=file_data,
                file_ext=file_ext,
                kb_id=kb_id,
                doc_id=document_id,
                tenant_id=tenant_id,
                title=title,
                source_type=source_type,
                source_url=source_url,
            )
            total_chars = sum(len(doc.page_content or "") for doc in documents)
            parser_type = documents[0].metadata.get("parser_type", "unknown") if documents else "empty"
            logger.info(
                f"[Stage2] 文档解析成功, parser={parser_type}, docs={len(documents)}, 文本长度={total_chars}"
            )
        except Exception as e:
            logger.error(f"[Stage2] 解析失败: {e}")
            update_document_status(document_id, "failed", error_message=f"解析失败: {str(e)}")
            raise
        
        # ===== Stage4: 清洗文本 =====
        logger.info(f"[Stage3] 开始清洗文本")
        cleaner = CleanProcessor(process_rule={"remove_extra_spaces": True, "remove_urls_emails": False})
        documents = cleaner.clean(documents)
        logger.info(f"[Stage3] 清洗完成")
        
        update_document_status(document_id, "indexing", progress=30)
        
        # ===== Stage5: 语义分块 =====
        logger.info(f"[Stage4] 开始语义分块, 文件类型: {file_ext}")
        try:
            chunks_data = build_chunks_data(documents, file_ext)
            chunk_type = chunks_data[0].get("chunk_type", "unknown") if chunks_data else "empty"
            logger.info(f"[Stage4] 分块完成, strategy={chunk_type}, 共 {len(chunks_data)} 个chunks")
        except Exception as e:
            logger.error(f"[Stage4] 分块失败: {e}")
            update_document_status(document_id, "failed", error_message=f"分块失败: {str(e)}")
            raise
        
        if not chunks_data:
            logger.warning(f"[Stage4] 无有效文本块")
            update_document_status(document_id, "failed", error_message="无有效文本内容")
            return {"status": "error", "message": "No valid text chunks"}
        
        update_document_status(document_id, "indexing", progress=80, chunk_count=len(chunks_data))
        
        # ===== 向量化（索引期不再逐 chunk 调 LLM） =====
        # 重构说明: 摘要与 HyDE 不再在索引期生成。
        # - 摘要仅用于展示,可后续异步/廉价生成,不阻塞索引;
        # - HyDE 改为「查询时」对用户 query 生成假设文档再检索(见 Retriever.search)。
        # 因此对原文直接 embedding,索引期零 LLM 调用(125 chunk 由 3~5 分钟降到秒级)。
        for c in chunks_data:
            c.setdefault("summary", "")
            c.setdefault("hypothetical_questions", [])
            c.setdefault("chunk_type", "recursive")

        logger.info(f"[Index] 开始向量化, 向量化目标: 原文内容, chunks数量: {len(chunks_data)}")
        
        texts_to_embed = [build_embedding_text(c) for c in chunks_data]
        
        try:
            vectors = embed_texts(texts_to_embed)
            logger.info(f"[Stage6] 向量化成功, 向量维度: {len(vectors[0]) if vectors else 0}")
        except Exception as e:
            logger.error(f"[Stage6] 向量化失败: {e}")
            update_document_status(document_id, "failed", error_message=f"向量化失败: {str(e)}")
            raise
        
        # ===== Stage8: 批量入库（事务：chunks + 状态） =====
        logger.info(f"[Stage7] 开始批量入库, chunks数量: {len(chunks_data)}")
        try:
            batch_insert_chunks(document_id, chunks_data, vectors)
            logger.info(f"[Stage7] 批量入库成功，状态已更新为ready")
        except Exception as e:
            logger.error(f"[Stage7] 入库失败: {e}")
            update_document_status(document_id, "failed", error_message=f"入库失败: {str(e)}")
            raise
        
        # ===== Stage9: 图片关联 =====
        logger.info(f"[Stage8] 开始图片关联")
        try:
            from core.rag.image_binder import ImageBinder
            binder = ImageBinder()
            binder.bind_images_to_chunks(
                doc_id=document_id,
                tenant_id=tenant_id,
                chunks_data=chunks_data,
            )
            binder.close()
            logger.info(f"[Stage8] 图片关联完成")
        except Exception as e:
            logger.warning(f"[Stage8] 图片关联失败: {e}")
        
        # ===== 完成 =====
        # 注意：状态已在 batch_insert_chunks 事务中更新为 ready
        cost_time = time.time() - start_time
        logger.info(f"[完成] 文档索引全部完成! doc_id={document_id}, chunks={len(chunks_data)}, 耗时={cost_time:.1f}秒")
        
        return {
            "status": "success",
            "document_id": document_id,
            "chunk_count": len(chunks_data),
            "cost_time": cost_time,
        }
        
    except Exception as e:
        error_msg = str(e)
        logger.error(f"[Celery] 索引失败: doc_id={document_id}, error={error_msg}")
        
        # 检查是否是可重试错误
        if self.request.retries < MAX_RETRIES:
            retry_delay = RETRY_DELAYS[self.request.retries]
            logger.info(f"[Celery] 第{self.request.retries+1}次重试, 等待{retry_delay}秒")
            raise self.retry(exc=e, countdown=retry_delay)
        else:
            update_document_status(document_id, "failed", error_message=error_msg)
            return {"status": "error", "message": error_msg}
    
    finally:
        engine.dispose()


@celery_app.task
def check_index_status(document_id: str) -> dict:
    """检查文档索引状态"""
    engine = create_engine(settings.sync_database_url)
    with Session(engine) as db:
        result = db.execute(
            text("SELECT status, progress, chunk_count, error_message, indexed_at FROM documents WHERE id=:id"),
            {"id": document_id}
        )
        row = result.mappings().first()
        
        if not row:
            return {"status": "not_found"}
        
        return {
            "status": row["status"],
            "progress": row["progress"],
            "chunk_count": row["chunk_count"],
            "error_message": row["error_message"],
            "indexed_at": row["indexed_at"].isoformat() if row["indexed_at"] else None,
        }
